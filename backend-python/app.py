from __future__ import annotations

import json
import shutil
import tempfile
from dataclasses import dataclass
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.document import Document as DocumentObject
from docx.table import _Cell, _Row, Table
from docx.text.paragraph import Paragraph
from flask import Flask, jsonify, request, send_file
from flask_cors import CORS
from werkzeug.utils import secure_filename

app = Flask(__name__)
CORS(app)

UPLOAD_DIR = Path(__file__).parent / "uploads"
UPLOAD_DIR.mkdir(exist_ok=True)


@dataclass
class LocatedComment:
    start: int
    end: int
    text: str
    comment: str
    author: str
    initials: str


@app.post("/api/comments/generate")
def generate_comments_docx():
    uploaded_file = request.files.get("file")
    if uploaded_file is None:
        return jsonify({"message": "请上传 DOCX 文件"}), 400

    comments_payload = request.form.get("comments", "[]")
    try:
        comments = json.loads(comments_payload)
    except json.JSONDecodeError:
        return jsonify({"message": "批注数据不是合法 JSON"}), 400

    if not isinstance(comments, list) or len(comments) == 0:
        return jsonify({"message": "至少需要一条批注"}), 400

    original_filename = uploaded_file.filename or "document.docx"
    if not original_filename.lower().endswith(".docx"):
        return jsonify({"message": "仅支持 .docx 文件"}), 400

    safe_name = secure_filename(original_filename)
    safe_stem = Path(safe_name).stem or "document"
    filename = f"{safe_stem}.docx"

    work_dir = Path(tempfile.mkdtemp(dir=UPLOAD_DIR))
    input_path = work_dir / filename
    output_path = work_dir / append_suffix(filename, "-comments")
    uploaded_file.save(input_path)

    document = Document(input_path)
    applied, skipped = apply_comments(document, comments)
    document.save(output_path)

    response = send_file(
        output_path,
        as_attachment=True,
        download_name=output_path.name,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    response.call_on_close(lambda: shutil.rmtree(work_dir, ignore_errors=True))
    response.headers["X-Comments-Applied"] = str(applied)
    response.headers["X-Comments-Skipped"] = str(skipped)
    return response


def apply_comments(document: DocumentObject, comments: list[dict]) -> tuple[int, int]:
    pending = [normalize_comment(comment) for comment in comments]
    paragraph_map: dict[Paragraph, list[LocatedComment]] = {}
    skipped = 0

    paragraphs = list(iter_paragraphs(document))
    for comment in pending:
        location = locate_comment(paragraphs, comment)
        if location is None:
            skipped += 1
            continue

        paragraph, located = location
        paragraph_map.setdefault(paragraph, []).append(located)

    applied = 0
    for paragraph, located_comments in paragraph_map.items():
        applied += rebuild_paragraph_with_comments(document, paragraph, located_comments)

    return applied, len(pending) - applied


def normalize_comment(comment: dict) -> dict:
    return {
        "selectedText": normalize_text(str(comment.get("selectedText", ""))),
        "occurrence": int(comment.get("occurrence", 0) or 0),
        "comment": str(comment.get("comment", "")).strip(),
        "author": str(comment.get("author", "Reviewer")).strip() or "Reviewer",
        "initials": str(comment.get("initials", "RV")).strip() or "RV",
    }


def locate_comment(
    paragraphs: list[Paragraph],
    comment: dict,
) -> tuple[Paragraph, LocatedComment] | None:
    selected_text = comment["selectedText"]
    if not selected_text or not comment["comment"]:
        return None

    seen = 0
    for paragraph in paragraphs:
        paragraph_text = normalize_text(paragraph.text)
        search_from = 0

        while True:
            start = paragraph_text.find(selected_text, search_from)
            if start == -1:
                break

            if seen == comment["occurrence"]:
                return (
                    paragraph,
                    LocatedComment(
                        start=start,
                        end=start + len(selected_text),
                        text=selected_text,
                        comment=comment["comment"],
                        author=comment["author"],
                        initials=comment["initials"],
                    ),
                )

            seen += 1
            search_from = start + len(selected_text)

    return None


def rebuild_paragraph_with_comments(
    document: DocumentObject,
    paragraph: Paragraph,
    located_comments: list[LocatedComment],
) -> int:
    text = normalize_text(paragraph.text)
    ordered = sorted(located_comments, key=lambda item: item.start)
    non_overlapping: list[LocatedComment] = []
    cursor = 0

    for item in ordered:
        if item.start < cursor:
            continue
        non_overlapping.append(item)
        cursor = item.end

    paragraph.clear()
    position = 0
    runs_for_comments: list[tuple[LocatedComment, object]] = []

    for item in non_overlapping:
        if item.start > position:
            paragraph.add_run(text[position:item.start])

        comment_run = paragraph.add_run(text[item.start:item.end])
        runs_for_comments.append((item, comment_run))
        position = item.end

    if position < len(text):
        paragraph.add_run(text[position:])

    for item, run in runs_for_comments:
        document.add_comment(
            runs=[run],
            text=item.comment,
            author=item.author,
            initials=item.initials,
        )

    return len(runs_for_comments)


def iter_paragraphs(parent: DocumentObject | _Cell | _Row | Table) -> Iterable[Paragraph]:
    if isinstance(parent, DocumentObject):
        for paragraph in parent.paragraphs:
            yield paragraph
        for table in parent.tables:
            yield from iter_paragraphs(table)
    elif isinstance(parent, Table):
        for row in parent.rows:
            yield from iter_paragraphs(row)
    elif isinstance(parent, _Row):
        for cell in parent.cells:
            yield from iter_paragraphs(cell)
    elif isinstance(parent, _Cell):
        for paragraph in parent.paragraphs:
            yield paragraph
        for table in parent.tables:
            yield from iter_paragraphs(table)


def normalize_text(value: str) -> str:
    return " ".join(value.split())


def append_suffix(filename: str, suffix: str) -> str:
    path = Path(filename)
    return f"{path.stem}{suffix}{path.suffix}"


@app.get("/api/health")
def health():
    return jsonify({"ok": True})


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=5000, debug=True)
